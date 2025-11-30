import os
import re
from io import BytesIO
import streamlit as st


_md_link = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
_md_bold = re.compile(r"(\*\*|__)(.*?)\1")
_md_italic = re.compile(r"(\*|_)(.*?)\1")


def markdown_to_plain_text(md: str) -> str:
    """Chuyển markdown -> text để PDF/DOCX không bị ký tự markdown."""
    if not md:
        return ""

    text = md.strip()

    # remove code fences
    text = re.sub(r"```.*?```", "", text, flags=re.DOTALL)

    # inline code
    text = re.sub(r"`([^`]+)`", r"\1", text)

    # headings -> remove #'s
    text = re.sub(r"^(#{1,6})\s+", "", text, flags=re.MULTILINE)

    # blockquote
    text = re.sub(r"^\s*>\s+", "", text, flags=re.MULTILINE)

    # links / emphasis
    text = _md_link.sub(r"\1 (\2)", text)
    text = _md_bold.sub(r"\2", text)
    text = _md_italic.sub(r"\2", text)

    # normalize bullets
    text = re.sub(r"^\s*[-*]\s+", "- ", text, flags=re.MULTILINE)

    # trim excessive blank lines
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    return text


def pdf_bytes_from_text(text: str) -> bytes:
    """Xuất PDF chỉ chứa nội dung AI (không thêm header/title của app)."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except Exception as e:
        raise RuntimeError(f"Thiếu thư viện reportlab: {e}")

    # Try register a Unicode font (để không bị ô vuông tiếng Việt)
    font_name = "Helvetica"
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSansCondensed.ttf",
        "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
    ]
    for fp in candidates:
        if os.path.exists(fp):
            try:
                pdfmetrics.registerFont(TTFont("UnicodeFont", fp))
                font_name = "UnicodeFont"
                break
            except Exception:
                pass

    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4

    margin_x = 40
    margin_y = 50
    y = height - margin_y

    c.setFont(font_name, 11)

    def wrap_line(line: str, max_width: float):
        words = line.split(" ")
        out, cur = [], ""
        for w in words:
            test = (cur + " " + w).strip()
            if pdfmetrics.stringWidth(test, font_name, 11) <= max_width:
                cur = test
            else:
                if cur:
                    out.append(cur)
                cur = w
        if cur:
            out.append(cur)
        return out

    max_w = width - 2 * margin_x
    for raw in (text or "").splitlines():
        line = raw.rstrip()
        if not line:
            y -= 10
            if y < margin_y:
                c.showPage()
                y = height - margin_y
                c.setFont(font_name, 11)
            continue

        for wl in wrap_line(line, max_w):
            c.drawString(margin_x, y, wl)
            y -= 14
            if y < margin_y:
                c.showPage()
                y = height - margin_y
                c.setFont(font_name, 11)

    c.showPage()
    c.save()
    return buf.getvalue()


def docx_bytes_from_text(text: str) -> bytes:
    """Xuất DOCX chỉ chứa nội dung AI (không thêm header/title của app)."""
    try:
        from docx import Document
    except Exception as e:
        raise RuntimeError(f"Thiếu thư viện python-docx: {e}")

    doc = Document()
    for line in (text or "").splitlines():
        line = line.rstrip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def safe_filename(name: str) -> str:
    name = (name or "Bao_cao").strip()
    name = re.sub(r"[^\w\-\.]+", "_", name, flags=re.UNICODE)
    return name[:120]
